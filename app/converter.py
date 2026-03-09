import io
import re
from docx import Document
from lxml import etree
from typing import List, Dict, Any, Tuple, Optional

# Quality commitment boilerplate with link
QUALITY_COMMITMENT_HTML = (
    '<div class="quality-commitment" data-skip-algolia="1">'
    'All content published on Kenhub is reviewed by medical and anatomy experts. '
    'The information we provide is grounded on academic literature and peer-reviewed research. '
    '<strong>Kenhub does not provide medical advice.</strong> '
    'You can learn more about our content creation and review standards by reading our '
    '<a href="/en/quality">content quality guidelines</a>.'
    '</div>'
)

# Regex to detect URLs in text
URL_RE = re.compile(r'(https?://[^\s,)<>\]]+)')

# Patterns that should NOT trigger the [Caption]Text special marker
BRACKET_SKIP_WORDS = re.compile(
    r'^\[(Updated|Internet|cited|Accessed|Retrieved|Available)\b',
    re.IGNORECASE
)


def linkify_urls(text: str) -> str:
    """Convert URLs in text to <a> tags with truncated display text."""
    def _replace(m):
        url = m.group(1)
        # Remove trailing punctuation that got captured
        while url and url[-1] in '.;':
            url = url[:-1]
        display = url[:30] + '...' if len(url) > 30 else url
        return f'<a href="{url}">{display}</a>'
    return URL_RE.sub(_replace, text)


def strip_bold_tags(html: str) -> str:
    """Remove <strong> tags from HTML, keeping inner text."""
    return re.sub(r'</?strong>', '', html)


def convert_docx_to_html(content: bytes) -> Dict[str, Any]:
    """
    Main conversion function that takes DOCX bytes and returns HTML + metadata
    """
    # Parse the DOCX
    doc = Document(io.BytesIO(content))
    
    # Extract metadata and determine content type
    metadata, content_type = extract_metadata_and_type(doc)
    
    # Process the document content in document order (paragraphs AND tables interleaved)
    html_parts = []
    warnings = []
    
    # Track list state
    current_list = None
    current_list_id = None
    
    # Track references/sources collection
    in_references_section = False   # After "Reference"/"References" heading
    in_sources_section = False      # After "Sources" heading
    collected_references = []       # Reference <li> items from Reference heading
    sources_parts = []              # Parts inside Sources section
    sources_refs = []               # Reference <li> items from Sources section
    sources_has_references_label = False
    sources_boilerplate_found = False
    
    # Build lookup maps from XML elements to python-docx objects
    WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    
    # Determine which paragraph index to start from (skip metadata)
    start_idx = skip_metadata_paragraphs(doc.paragraphs)
    skip_elements = set()
    for i in range(start_idx):
        skip_elements.add(id(doc.paragraphs[i]._element))
    
    # Iterate body elements in document order
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        
        if tag == 'p' and element in para_map:
            if id(element) in skip_elements:
                continue
            paragraph = para_map[element]
            
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # ── Detect section headings ──
            is_heading = style_name in ("Heading 1", "Heading 2", "Heading 3", "Title")
            
            if is_heading:
                # Check for Reference/References heading
                if re.match(r'^references?$', text, re.IGNORECASE):
                    # Close any open list
                    if current_list:
                        html_parts.append(f"</{current_list}>")
                        current_list = None
                        current_list_id = None
                    in_references_section = True
                    in_sources_section = False
                    continue
                
                # Check for Sources heading
                if re.match(r'^sources$', text, re.IGNORECASE):
                    # Close any open list
                    if current_list:
                        html_parts.append(f"</{current_list}>")
                        current_list = None
                        current_list_id = None
                    in_references_section = False
                    in_sources_section = True
                    continue
                
                # Any other heading exits reference/sources collection
                if in_references_section or in_sources_section:
                    _flush_sources(html_parts, collected_references, sources_parts,
                                   sources_refs, sources_has_references_label,
                                   sources_boilerplate_found)
                    in_references_section = False
                    in_sources_section = False
                    collected_references = []
                    sources_parts = []
                    sources_refs = []
                    sources_has_references_label = False
                    sources_boilerplate_found = False
            
            # ── Collecting references from "Reference" section ──
            if in_references_section:
                ref_text = _extract_reference_text(paragraph)
                if ref_text:
                    ref_html = linkify_urls(ref_text)
                    collected_references.append(f"  <li>{ref_html}</li>")
                continue
            
            # ── Collecting items in "Sources" section ──
            if in_sources_section:
                # Detect boilerplate paragraph
                if 'all content published on kenhub' in text.lower():
                    sources_boilerplate_found = True
                    continue
                
                # Detect "References:" label
                if re.match(r'^references\s*:?\s*$', text, re.IGNORECASE):
                    sources_has_references_label = True
                    continue
                
                # Everything else is a reference item
                ref_text = _extract_reference_text(paragraph)
                if ref_text:
                    ref_html = linkify_urls(ref_text)
                    sources_refs.append(f"  <li>{ref_html}</li>")
                continue
            
            # ── Normal processing (not in references/sources) ──
            
            # Check for special markers — but not for reference-like text
            if not BRACKET_SKIP_WORDS.search(text):
                special_html, special_warnings = process_special_markers(text)
                if special_html:
                    if current_list:
                        html_parts.append(f"</{current_list}>")
                        current_list = None
                        current_list_id = None
                    html_parts.append(special_html)
                    warnings.extend(special_warnings)
                    continue
            
            # Check if this is a bullet character list item (● or •)
            bullet_text = _get_bullet_text(paragraph)
            
            # Check if this is a numPr list item
            list_info = get_list_info(paragraph)
            
            if bullet_text is not None:
                # Bullet character detected — treat as unordered list item
                if current_list != "ul" or current_list_id != "__bullet__":
                    if current_list:
                        html_parts.append(f"</{current_list}>")
                    current_list = "ul"
                    current_list_id = "__bullet__"
                    html_parts.append("<ul>")
                item_html = _process_bullet_paragraph_content(paragraph)
                html_parts.append(f"  <li>{item_html}</li>")
            elif list_info:
                list_type, list_id = list_info
                if current_list_id != list_id:
                    if current_list:
                        html_parts.append(f"</{current_list}>")
                    current_list = list_type
                    current_list_id = list_id
                    html_parts.append(f"<{list_type}>")
                item_html = process_paragraph_content(paragraph)
                html_parts.append(f"  <li>{item_html}</li>")
            else:
                if current_list:
                    html_parts.append(f"</{current_list}>")
                    current_list = None
                    current_list_id = None
                para_html = process_regular_paragraph(paragraph)
                if para_html:
                    html_parts.append(para_html)
        
        elif tag == 'tbl' and element in table_map:
            # Close any open list before table
            if current_list:
                html_parts.append(f"</{current_list}>")
                current_list = None
                current_list_id = None
            table_html = process_table(table_map[element])
            html_parts.append(table_html)
    
    # Close any remaining open list
    if current_list:
        html_parts.append(f"</{current_list}>")
    
    # Flush any remaining sources/references section
    if in_references_section or in_sources_section:
        _flush_sources(html_parts, collected_references, sources_parts,
                       sources_refs, sources_has_references_label,
                       sources_boilerplate_found)
    
    # Apply content wrappers based on type
    final_html = apply_content_wrappers(html_parts, content_type, metadata)
    
    return {
        "html": final_html,
        "metadata": metadata,
        "warnings": warnings
    }


def _get_bullet_text(paragraph) -> Optional[str]:
    """If paragraph starts with ● or •, return text after stripping bullet. Else None."""
    text = paragraph.text
    if not text:
        return None
    stripped = text.lstrip()
    # Handle repeated bullet chars like "● ●" — strip all leading bullets and whitespace
    if stripped and stripped[0] in '●•':
        # Strip all leading bullet chars and whitespace
        cleaned = stripped.lstrip('●• \t')
        return cleaned
    return None


def _process_bullet_paragraph_content(paragraph) -> str:
    """Process runs in a bullet paragraph, stripping leading bullet chars from first run."""
    result_parts = []
    first_text_seen = False
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        
        if not first_text_seen:
            # Strip leading bullet chars and whitespace from first non-empty run
            text = text.lstrip('●• \t')
            if not text:
                continue
            first_text_seen = True
        
        # Apply formatting
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        
        result_parts.append(text)
    
    return "".join(result_parts)


def _extract_reference_text(paragraph) -> Optional[str]:
    """Extract plain reference text from a paragraph, stripping bullets and bold formatting."""
    text = paragraph.text.strip()
    if not text:
        return None
    
    # Strip leading bullet characters
    text = text.lstrip('●• \t')
    if not text:
        return None
    
    # Build text from runs but strip bold formatting
    parts = []
    first_text_seen = False
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        if not first_text_seen:
            run_text = run_text.lstrip('●• \t')
            if not run_text:
                continue
            first_text_seen = True
        
        # Apply italic but NOT bold (references should be plain text except italic for titles)
        if run.italic:
            run_text = f"<em>{run_text}</em>"
        
        parts.append(run_text)
    
    return "".join(parts) if parts else text


def _flush_sources(html_parts, collected_references, sources_parts,
                   sources_refs, sources_has_references_label,
                   sources_boilerplate_found):
    """Flush collected references and sources into the article-meta-content wrapper."""
    # Merge all reference items
    all_refs = collected_references + sources_refs
    
    if not all_refs and not sources_boilerplate_found:
        return
    
    parts = ['<div class="article-meta-content">']
    parts.append('<h2>Sources</h2>')
    
    # Always include quality commitment
    parts.append(QUALITY_COMMITMENT_HTML)
    
    # "References:" label if present in source doc
    if sources_has_references_label:
        parts.append('<p>References:</p>')
    
    # Reference list
    if all_refs:
        parts.append('<ul>')
        parts.extend(all_refs)
        parts.append('</ul>')
    
    parts.append('</div>')
    html_parts.extend(parts)


def extract_metadata_and_type(doc: Document) -> Tuple[Dict[str, str], str]:
    """Extract metadata from document and determine content type"""
    metadata = {
        "title": "",
        "description": "",
        "type": "article",
        "seo_title": "",
        "seo_description": ""
    }
    
    # Look for metadata in first few paragraphs
    for para in doc.paragraphs[:10]:
        text = para.text.strip()
        if text.startswith("Title:"):
            metadata["title"] = text[6:].strip()
        elif text.startswith("Description:"):
            metadata["description"] = text[12:].strip()
        elif text.startswith("SEO title:"):
            metadata["seo_title"] = text[10:].strip()
        elif text.startswith("SEO description:"):
            metadata["seo_description"] = text[16:].strip()
    
    # Determine content type - study_unit if has learning objectives
    content_type = "article"
    for para in doc.paragraphs:
        if "learning objective" in para.text.lower() or "after completing this study unit" in para.text.lower():
            content_type = "study_unit"
            break
    
    metadata["type"] = content_type
    return metadata, content_type

def skip_metadata_paragraphs(paragraphs: List) -> int:
    """Skip metadata paragraphs at the start and return the index to start processing"""
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if not (text.startswith("Title:") or text.startswith("Description:") or 
                text.startswith("SEO title:") or text.startswith("SEO description:") or
                text.startswith("Container:") or text.startswith("Position:")):
            return i
    return 0

def get_list_info(paragraph) -> Optional[Tuple[str, str]]:
    """Check if paragraph is a list item and return (list_type, list_id)"""
    try:
        # Access the XML to check for numPr
        p_xml = paragraph._element
        numPr = p_xml.xpath('.//w:numPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        if numPr:
            numId_elem = numPr[0].xpath('./w:numId', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            ilvl_elem = numPr[0].xpath('./w:ilvl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if numId_elem:
                num_id = numId_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                return ("ul", num_id)
    except:
        pass
    
    return None

def process_paragraph_content(paragraph) -> str:
    """Process the runs within a paragraph to handle formatting"""
    result_parts = []
    
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        # Apply formatting
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        
        result_parts.append(text)
    
    return "".join(result_parts)

def process_regular_paragraph(paragraph) -> str:
    """Process a regular paragraph (not list item)"""
    style_name = paragraph.style.name if paragraph.style else "Normal"
    text = paragraph.text.strip()
    
    if not text:
        return ""
    
    # Skip metadata lines
    if any(text.startswith(prefix) for prefix in ["Title:", "Description:", "SEO title:", "SEO description:", "Container:", "Position:"]):
        return ""
    
    # Handle headings
    if style_name == "Heading 1" or style_name == "Title":
        return ""  # Skip titles as per rules
    elif style_name == "Heading 2":
        return f"<h2>{text}</h2>"
    elif style_name == "Heading 3":
        return f"<h3>{text}</h3>"
    else:
        # Regular paragraph
        content = process_paragraph_content(paragraph)
        return f"<p>{content}</p>"

def process_special_markers(text: str) -> Tuple[str, List[str]]:
    """Process special markers like [Video:], [Table of Contents], etc."""
    warnings = []
    
    # Video markers
    video_match = re.search(r'\[Video:\s*([^\]]*)\]', text, re.IGNORECASE)
    if video_match:
        warnings.append(f"Video placeholder needs ID assignment: {video_match.group(1)}")
        return '<div class="embedded-video-container embedded-widget" data-controller="embedded_video" data-embedded_video-id-value="XXXX" isvisible="true"></div>', warnings
    
    # Table of Contents
    if re.search(r'\[Table of [Cc]ontents\]', text):
        return '<div class="article-table-of-contents" data-skip-algolia="1">Table of Contents</div>', warnings
    
    # Blue box
    blue_box_match = re.search(r'\[Blue box:\s*([^\]]*)\]', text, re.IGNORECASE)
    if blue_box_match:
        content = blue_box_match.group(1)
        return f'<div class="highlighted-box"><p>{content}</p></div>', warnings
    
    # Green highlight / Highlights gallery
    highlight_match = re.search(r'\[(Green highlight|Highlights gallery)[^\]]*\]\s*([^\n]*)', text, re.IGNORECASE)
    if highlight_match:
        slugs = highlight_match.group(2).strip()
        return f'<div class="image-gallery-container embedded-widget outset-left open" data-image-slugs="{slugs}"></div>', warnings
    
    # Gallery pattern
    gallery_match = re.search(r'Gallery:\s*([^\n]*)', text, re.IGNORECASE)
    if gallery_match:
        terms = gallery_match.group(1).strip()
        # Convert terms to slugs (simplified)
        slugs = ",".join(term.strip().lower().replace(" ", "-") for term in terms.split(","))
        return f'<div class="image-gallery-container embedded-widget open" data-image-slugs="{slugs}"></div>', warnings
    
    # Overview images [Caption]Text pattern
    # Exclude brackets with dates, "Updated", "Internet", "cited" etc.
    caption_match = re.search(r'\[([^\]]+)\](.+)', text)
    if caption_match:
        bracket_content = caption_match.group(1)
        # Skip if bracket content looks like a citation annotation
        if not re.match(r'(Updated|Internet|cited|Accessed|Retrieved|Available)\b', bracket_content, re.IGNORECASE):
            warnings.append(f"Overview image placeholder needs ID assignment: {bracket_content}")
            return '<!-- OVERVIEW IMAGE: Manual ID assignment needed -->', warnings
    
    # Content box links
    if '/en/study/' in text or '/en/custom-quizzes/' in text:
        link = text.strip()
        return f'<div class="contentbox-container" data-contentbox-link="{link}" data-skip-algolia="1">{link}</div>', warnings
    
    return "", warnings

def process_table(table) -> str:
    """Convert DOCX table to HTML"""
    rows_html = []
    caption = ""
    
    # Check if first row is caption (merged cells)
    first_row = table.rows[0]
    if len(first_row.cells) == 1 or (len(first_row.cells) > 1 and first_row.cells[0].text.strip()):
        # Treat first row as caption if it seems like one
        cell_text = first_row.cells[0].text.strip()
        if cell_text and not any(cell.text.strip() for cell in first_row.cells[1:]):
            caption = cell_text
            start_row = 1
        else:
            start_row = 0
    else:
        start_row = 0
    
    # Process data rows
    for row in table.rows[start_row:]:
        cells_html = []
        for cell in row.cells:
            # Process cell content with line breaks
            cell_parts = []
            for para in cell.paragraphs:
                if para.text.strip():
                    cell_parts.append(process_paragraph_content(para))
            
            cell_content = "<br>".join(cell_parts)
            cells_html.append(f"<td>{cell_content}</td>")
        
        rows_html.append(f"    <tr>\n      {chr(10).join(cells_html)}\n    </tr>")
    
    # Build table HTML
    table_parts = ['<table class="facts-table">']
    if caption:
        table_parts.append(f"  <caption>{caption}</caption>")
    table_parts.append("  <tbody>")
    table_parts.extend(rows_html)
    table_parts.append("  </tbody>")
    table_parts.append("</table>")
    
    return "\n".join(table_parts)

def apply_content_wrappers(html_parts: List[str], content_type: str, metadata: Dict) -> str:
    """Apply content type specific wrappers"""
    content = "\n".join(html_parts)
    
    if content_type == "study_unit":
        # Find learning objectives section
        obj_match = re.search(r'(<h[23]>.*?learning objectives?.*?</h[23]>)(.*?)(<h[23]|$)', content, re.IGNORECASE | re.DOTALL)
        if obj_match:
            before = content[:obj_match.start()]
            objectives_section = obj_match.group(1) + obj_match.group(2)
            rest = content[obj_match.end()-len(obj_match.group(3)):]
            
            wrapped_objectives = f'<div class="highlighted-box">\n{objectives_section}\n</div>'
            wrapped_rest = f'<div class="learning-path">\n{rest}\n</div>' if rest.strip() else ''
            
            return f'{before}\n{wrapped_objectives}\n{wrapped_rest}'
    
    return content
